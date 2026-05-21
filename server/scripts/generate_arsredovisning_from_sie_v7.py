#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ZERO = Decimal("0")


@dataclass
class SieData:
    org_number: str = ""
    company_name: str = ""
    current_start: str = ""
    current_end: str = ""
    prior_start: str = ""
    prior_end: str = ""
    ib: Dict[int, Dict[int, Decimal]] = field(default_factory=dict)
    ub: Dict[int, Dict[int, Decimal]] = field(default_factory=dict)
    res: Dict[int, Dict[int, Decimal]] = field(default_factory=dict)


@dataclass
class BuiltValues:
    paragraphs: Dict[str, str]
    tables: Dict[str, List[List[str]]]
    raw: Dict[str, Any]


@dataclass
class RoundedSide:
    line_items: Dict[str, int]
    total: int
    rounding_diff_from_individual: int


@dataclass
class RoundedStatement:
    line_items: Dict[str, int]
    total: int
    rounding_diff_from_individual: int


def read_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: Dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_sie(path: Path) -> SieData:
    text = path.read_text(encoding="cp437")
    data = SieData(ib={}, ub={}, res={})

    for line in text.splitlines():
        if line.startswith("#ORGNR "):
            data.org_number = line.split(maxsplit=1)[1].strip()
        elif line.startswith("#FNAMN "):
            match = re.search(r'"(.*)"', line)
            data.company_name = match.group(1) if match else line.split(maxsplit=1)[1].strip()
        elif line.startswith("#RAR 0 "):
            _, _, start, end = line.split(maxsplit=3)
            data.current_start = format_sie_date(start)
            data.current_end = format_sie_date(end)
        elif line.startswith("#RAR -1 "):
            _, _, start, end = line.split(maxsplit=3)
            data.prior_start = format_sie_date(start)
            data.prior_end = format_sie_date(end)
        elif line.startswith("#IB "):
            _, year, account, amount = line.split(maxsplit=3)
            data.ib.setdefault(int(year), {})[int(account)] = Decimal(amount)
        elif line.startswith("#UB "):
            _, year, account, amount = line.split(maxsplit=3)
            data.ub.setdefault(int(year), {})[int(account)] = Decimal(amount)
        elif line.startswith("#RES "):
            _, year, account, amount = line.split(maxsplit=3)
            data.res.setdefault(int(year), {})[int(account)] = Decimal(amount)

    return data


def format_sie_date(value: str) -> str:
    value = value.strip()
    if not re.fullmatch(r"\d{8}", value):
        return value
    return f"{value[:4]}-{value[4:6]}-{value[6:8]}"


def format_kr(amount: int | Decimal | str) -> str:
    if isinstance(amount, str):
        return amount
    if not isinstance(amount, Decimal):
        amount = Decimal(amount)
    negative = amount < 0
    value = abs(int(amount))
    formatted = f"{value:,}".replace(",", " ")
    return f"-{formatted}" if negative else formatted


def round_half_up_int(value: Decimal) -> int:
    return int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def trunc_towards_zero(value: Decimal) -> int:
    return int(value)


def round_int_by_mode(value: Decimal, mode: str) -> int:
    if mode == "half_up":
        return round_half_up_int(value)
    if mode == "truncate":
        return trunc_towards_zero(value)
    raise ValueError(f"Okänt avrundningsläge: {mode}")


def prompt_text(label: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    reply = input(f"{label}{suffix}: ").strip()
    return reply if reply else default



def prompt_int(label: str, default: int = 0) -> int:
    while True:
        raw = prompt_text(label, str(default))
        try:
            return int(raw.replace(" ", ""))
        except ValueError:
            print("Ange ett heltal.")



def prompt_decimal_text(label: str, default: str = "") -> str:
    return prompt_text(label, default)


def prompt_bool(label: str, default: bool = False) -> bool:
    default_text = "j" if default else "n"
    while True:
        raw = prompt_text(f"{label} (j/n)", default_text).strip().lower()
        if raw in {"j", "ja", "y", "yes", "1", "true"}:
            return True
        if raw in {"n", "nej", "no", "0", "false"}:
            return False
        print("Svara j eller n.")


def get_account(source: Dict[int, Dict[int, Decimal]], year: int, account: int) -> Decimal:
    return source.get(year, {}).get(account, ZERO)



def sum_accounts(source: Dict[int, Dict[int, Decimal]], year: int, accounts: List[int]) -> Decimal:
    return sum((get_account(source, year, acc) for acc in accounts), start=ZERO)



def sum_account_range(source: Dict[int, Dict[int, Decimal]], year: int, start_acc: int, end_acc: int) -> Decimal:
    year_data = source.get(year, {})
    return sum((amount for acc, amount in year_data.items() if start_acc <= acc <= end_acc), start=ZERO)



def to_positive(value: Decimal) -> Decimal:
    return abs(value)



def round_with_target(values: Dict[str, Decimal], mode: str, target_total: Optional[int] = None) -> Dict[str, Any]:
    """Round values to whole kronor while forcing the displayed sum to match target_total.

    The function starts from ordinary half-up rounding and only adjusts the minimum
    number of rows needed to hit the requested total, choosing the rows where the
    extra rounding error is smallest.
    """
    if target_total is None:
        target_total = round_int_by_mode(sum(values.values(), start=ZERO), mode)

    rounded = {name: round_int_by_mode(value, mode) for name, value in values.items()}
    current_total = sum(rounded.values())
    diff = target_total - current_total

    if diff == 0:
        return {
            "rounded": rounded,
            "target_total": target_total,
            "rounding_diff_from_individual": 0,
        }

    def adjustment_cost(name: str, step: int) -> Decimal:
        value = values[name]
        current = Decimal(rounded[name])
        new = Decimal(rounded[name] + step)
        return abs(new - value) - abs(current - value)

    names = list(values.keys())
    step = 1 if diff > 0 else -1
    for _ in range(abs(diff)):
        candidates = []
        for name in names:
            if step < 0 and rounded[name] <= 0 and values[name] >= 0:
                continue
            cost = adjustment_cost(name, step)
            frac = values[name] - Decimal(int(values[name]))
            candidates.append((cost, -abs(values[name]), -frac, name))
        if not candidates:
            raise ValueError("Kunde inte balansera avrundningen med givna värden.")
        candidates.sort()
        chosen = candidates[0][3]
        rounded[chosen] += step

    return {
        "rounded": rounded,
        "target_total": target_total,
        "rounding_diff_from_individual": diff,
    }



def round_balance_side(values: Dict[str, Decimal], shared_target_total: int, mode: str) -> RoundedSide:
    rounded = round_with_target(values, mode=mode, target_total=shared_target_total)
    return RoundedSide(
        line_items=rounded["rounded"],
        total=rounded["target_total"],
        rounding_diff_from_individual=rounded["rounding_diff_from_individual"],
    )



def round_result_statement(values: Dict[str, Decimal], mode: str) -> RoundedStatement:
    rounded = round_with_target(values, mode=mode)
    return RoundedStatement(
        line_items=rounded["rounded"],
        total=rounded["target_total"],
        rounding_diff_from_individual=rounded["rounding_diff_from_individual"],
    )



def default_manual_from_sie(sie: SieData) -> Dict[str, Any]:
    current_year = int(sie.current_start[:4]) if sie.current_start[:4].isdigit() else 2024
    prior_year = int(sie.prior_start[:4]) if sie.prior_start[:4].isdigit() else current_year - 1
    older_1 = current_year - 2
    older_2 = current_year - 3
    return {
        "company_name": sie.company_name,
        "org_number": sie.org_number,
        "report_start": sie.current_start,
        "report_end": sie.current_end,
        "prior_start": sie.prior_start,
        "prior_end": sie.prior_end,
        "board_city": "",
        "adoption_date": "",
        "faststallelse_signatory": "",
        "faststallelse_title": "Styrelseledamot",
        "business_description": "",
        "registered_seat": "",
        "avg_employees": "0,0",
        "utdelning": 0,
        "report_mode": "paper",
        "submission_date": "",
        "document_date": "",
        "has_auditor": False,
        "significant_events_during_year": "",
        "significant_events_after_year_end": "",
        "has_own_shares": False,
        "own_shares_text": "",
        "net_sales_variation_comment": "",
        "company_is_public": False,
        "parent_in_larger_group": False,
        "parent_in_smaller_group_prepares_consolidated": False,
        "foreign_branch": False,
        "share_based_payments": False,
        "compound_instruments": False,
        "crypto_assets": False,
        "k2_previous_year": True,
        "deferred_tax_liability_significant": False,
        "buildings_generate_75_pct_turnover": False,
        "headcount_over_3_two_years": False,
        "balance_over_1_5m_two_years": False,
        "net_sales_over_3m_two_years": False,
        "tax_rate_soliditet": 0.206,
        "rounding_mode": "truncate",
        "board_members": [
            {"name": "", "title": "", "date": ""},
            {"name": "", "title": "", "date": ""},
        ],
        "flerarsoversikt": {
            str(older_1): {
                "start": f"{older_1}-01-01",
                "end": f"{older_1}-12-31",
                "nettoomsattning": 0,
                "resultat_efter_finansiella_poster": 0,
                "soliditet": 0,
            },
            str(older_2): {
                "start": f"{older_2}-01-01",
                "end": f"{older_2}-12-31",
                "nettoomsattning": 0,
                "resultat_efter_finansiella_poster": 0,
                "soliditet": 0,
            },
        },
        "overrides": {},
    }



def merge_manual(base: Dict[str, Any], extra: Dict[str, Any]) -> Dict[str, Any]:
    result = deepcopy(base)
    for key, value in extra.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = merge_manual(result[key], value)
        else:
            result[key] = deepcopy(value)
    return result



def collect_manual_data(sie: SieData, manual: Dict[str, Any]) -> Dict[str, Any]:
    data = merge_manual(default_manual_from_sie(sie), manual)
    print("\nFyll i de uppgifter som inte finns i SIE-filen. Tryck Enter för att behålla standardvärdet.\n")

    data["company_name"] = prompt_text("Bolagsnamn", str(data.get("company_name", "")))
    data["org_number"] = prompt_text("Organisationsnummer", str(data.get("org_number", "")))
    data["report_start"] = prompt_text("Rapportperiod start", str(data.get("report_start", "")))
    data["report_end"] = prompt_text("Rapportperiod slut", str(data.get("report_end", "")))
    data["prior_start"] = prompt_text("Föregående år start", str(data.get("prior_start", "")))
    data["prior_end"] = prompt_text("Föregående år slut", str(data.get("prior_end", "")))
    data["business_description"] = prompt_text("Verksamhetsbeskrivning", str(data.get("business_description", "")))
    data["registered_seat"] = prompt_text("Säte", str(data.get("registered_seat", "")))
    data["board_city"] = prompt_text("Ort för underskrift", str(data.get("board_city", "")))
    data["adoption_date"] = prompt_text("Fastställelsedatum / stämmodatum (YYYY-MM-DD)", str(data.get("adoption_date", "")))
    data["faststallelse_signatory"] = prompt_text("Namn i fastställelseintyget", str(data.get("faststallelse_signatory", "")))
    data["faststallelse_title"] = prompt_text("Roll i fastställelseintyget", str(data.get("faststallelse_title", "Styrelseledamot")))
    data["avg_employees"] = prompt_decimal_text("Medelantal anställda", str(data.get("avg_employees", "0,0")))
    data["utdelning"] = prompt_int("Utdelning", int(data.get("utdelning", 0)))
    while True:
        report_mode = prompt_text("Inlämningssätt (paper eller digital)", str(data.get("report_mode", "paper"))).strip().lower()
        if report_mode in {"paper", "digital"}:
            data["report_mode"] = report_mode
            break
        print("Ange paper eller digital.")
    data["document_date"] = prompt_text("Dateringsdatum för årsredovisningen (YYYY-MM-DD)", str(data.get("document_date", data.get("adoption_date", ""))))
    data["submission_date"] = prompt_text("Planerat inlämningsdatum till Bolagsverket (YYYY-MM-DD)", str(data.get("submission_date", "")))
    data["has_auditor"] = prompt_bool("Har bolaget revisor", bool(data.get("has_auditor", False)))
    data["significant_events_during_year"] = prompt_text("Väsentliga händelser under räkenskapsåret (tomt om inga)", str(data.get("significant_events_during_year", "")))
    data["significant_events_after_year_end"] = prompt_text("Väsentliga händelser efter räkenskapsårets slut (tomt om inga)", str(data.get("significant_events_after_year_end", "")))
    data["has_own_shares"] = prompt_bool("Har bolaget egna aktier som kräver upplysning", bool(data.get("has_own_shares", False)))
    if data["has_own_shares"]:
        data["own_shares_text"] = prompt_text("Text för upplysning om egna aktier", str(data.get("own_shares_text", "")))
    else:
        data["own_shares_text"] = ""
    data["net_sales_variation_comment"] = prompt_text("Kommentar om nettoomsättningen varierar mer än 30 procent (tomt om ej aktuellt)", str(data.get("net_sales_variation_comment", "")))

    print("\nK2-kontroll enligt BFNAR 2016:10 / ändringar för senare räkenskapsår.\n")
    data["company_is_public"] = prompt_bool("Är bolaget publikt", bool(data.get("company_is_public", False)))
    data["parent_in_larger_group"] = prompt_bool("Är bolaget moderföretag i större koncern", bool(data.get("parent_in_larger_group", False)))
    data["parent_in_smaller_group_prepares_consolidated"] = prompt_bool(
        "Är bolaget moderföretag i mindre koncern och upprättar koncernredovisning",
        bool(data.get("parent_in_smaller_group_prepares_consolidated", False)),
    )
    data["foreign_branch"] = prompt_bool("Har bolaget filial i utlandet under året", bool(data.get("foreign_branch", False)))
    data["share_based_payments"] = prompt_bool("Har bolaget aktierelaterade ersättningar", bool(data.get("share_based_payments", False)))
    data["compound_instruments"] = prompt_bool("Har bolaget konvertibler eller liknande sammansatta finansiella instrument", bool(data.get("compound_instruments", False)))
    data["crypto_assets"] = prompt_bool("Har bolaget kryptotillgångar utöver enstaka betalningar", bool(data.get("crypto_assets", False)))
    data["k2_previous_year"] = prompt_bool("Tillämpade bolaget K2 föregående räkenskapsår", bool(data.get("k2_previous_year", True)))
    data["deferred_tax_liability_significant"] = prompt_bool("Har bolaget väsentlig uppskjuten skatteskuld", bool(data.get("deferred_tax_liability_significant", False)))
    data["buildings_generate_75_pct_turnover"] = prompt_bool("Har bolaget byggnader som genererar minst 75 procent av nettoomsättningen", bool(data.get("buildings_generate_75_pct_turnover", False)))
    data["headcount_over_3_two_years"] = prompt_bool("Har medelantalet anställda överstigit 3 båda de två senaste åren", bool(data.get("headcount_over_3_two_years", False)))
    data["balance_over_1_5m_two_years"] = prompt_bool("Har balansomslutningen överstigit 1,5 mkr båda de två senaste åren", bool(data.get("balance_over_1_5m_two_years", False)))
    data["net_sales_over_3m_two_years"] = prompt_bool("Har nettoomsättningen överstigit 3 mkr båda de två senaste åren", bool(data.get("net_sales_over_3m_two_years", False)))

    while True:
        rounding_mode = prompt_text("Avrundningsläge (truncate eller half_up)", str(data.get("rounding_mode", "truncate"))).strip()
        if rounding_mode in {"truncate", "half_up"}:
            data["rounding_mode"] = rounding_mode
            break
        print("Ange truncate eller half_up.")

    rate_default = str(data.get("tax_rate_soliditet", 0.206))
    while True:
        raw_rate = prompt_text("Skattesats för soliditet (t.ex. 0.206)", rate_default).replace(",", ".")
        try:
            data["tax_rate_soliditet"] = float(raw_rate)
            break
        except ValueError:
            print("Ange ett decimaltal, t.ex. 0.206")

    current_year = int(data["report_start"][:4]) if str(data["report_start"])[:4].isdigit() else 2024
    years = [str(current_year - 2), str(current_year - 3)]
    data.setdefault("flerarsoversikt", {})
    for year in years:
        row = data["flerarsoversikt"].setdefault(year, {})
        row["start"] = prompt_text(f"Flerårsöversikt {year} start", str(row.get("start", f"{year}-01-01")))
        row["end"] = prompt_text(f"Flerårsöversikt {year} slut", str(row.get("end", f"{year}-12-31")))
        row["nettoomsattning"] = prompt_int(f"Flerårsöversikt {year} nettoomsättning", int(row.get("nettoomsattning", 0)))
        row["resultat_efter_finansiella_poster"] = prompt_int(
            f"Flerårsöversikt {year} resultat efter finansiella poster",
            int(row.get("resultat_efter_finansiella_poster", 0)),
        )
        row["soliditet"] = prompt_int(f"Flerårsöversikt {year} soliditet (%)", int(row.get("soliditet", 0)))

    members: List[Dict[str, Any]] = list(data.get("board_members", []))
    default_count = max(2, len(members)) if members else 2
    member_count = prompt_int("Antal underskrifter i dokumentet", default_count)
    while len(members) < member_count:
        members.append({"name": "", "title": "", "date": data.get("adoption_date", "")})

    for idx in range(member_count):
        members[idx]["name"] = prompt_text(f"Underskrift {idx + 1} namn", str(members[idx].get("name", "")))
        members[idx]["title"] = prompt_text(f"Underskrift {idx + 1} roll", str(members[idx].get("title", "")))
        members[idx]["date"] = prompt_text(
            f"Underskrift {idx + 1} datum (YYYY-MM-DD)",
            str(members[idx].get("date", data.get("adoption_date", ""))),
        )
    data["board_members"] = members[:member_count]

    data.setdefault("overrides", {})
    return data



def ensure_manual_data(sie: SieData, manual: Dict[str, Any], interactive: bool) -> Dict[str, Any]:
    merged = merge_manual(default_manual_from_sie(sie), manual)

    required_paths = [
        "business_description",
        "registered_seat",
        "board_city",
        "adoption_date",
        "faststallelse_signatory",
        "faststallelse_title",
        "avg_employees",
    ]

    current_year = int(merged["report_start"][:4]) if str(merged["report_start"])[:4].isdigit() else 2024
    for year in (str(current_year - 2), str(current_year - 3)):
        required_paths.extend([
            f"flerarsoversikt.{year}.nettoomsattning",
            f"flerarsoversikt.{year}.resultat_efter_finansiella_poster",
            f"flerarsoversikt.{year}.soliditet",
        ])

    board_members = merged.get("board_members", [])
    if not board_members:
        required_paths.append("board_members")
    else:
        for idx, _member in enumerate(board_members):
            required_paths.extend([f"board_members.{idx}.name", f"board_members.{idx}.title"])

    if merged.get("has_own_shares"):
        required_paths.append("own_shares_text")

    missing = []
    for path in required_paths:
        value = get_nested_value(merged, path)
        if value in (None, ""):
            missing.append(path)

    if interactive or (missing and sys.stdin.isatty()):
        return collect_manual_data(sie, merged)

    if missing:
        missing_text = "\n".join(f"- {item}" for item in missing)
        raise ValueError(
            "Den manuella JSON-filen saknar obligatoriska uppgifter. Kör med --interactive eller fyll i följande fält:\n"
            + missing_text
        )
    return merged



def get_nested_value(data: Any, path: str) -> Any:
    current = data
    for part in path.split("."):
        if isinstance(current, list):
            try:
                current = current[int(part)]
            except (ValueError, IndexError):
                return None
        elif isinstance(current, dict):
            if part not in current:
                return None
            current = current[part]
        else:
            return None
    return current



def build_values(sie: SieData, manual: Dict[str, Any]) -> BuiltValues:
    current_year = sie.current_start[:4]
    prior_year = sie.prior_start[:4]
    report_year = int(manual.get("report_start", sie.current_start)[:4]) if str(manual.get("report_start", sie.current_start))[:4].isdigit() else int(current_year)
    older = manual.get("flerarsoversikt", {})
    tax_rate = Decimal(str(manual.get("tax_rate_soliditet", 0.206)))
    rounding_mode = str(manual.get("rounding_mode", "truncate"))

    # Resultaträkning råvärden (tecken som ska visas i rapporten)
    result_current_raw = {
        "net_sales": sum_account_range(sie.res, 0, 3000, 3799),
        "external": -sum_accounts(sie.res, 0, [6570]),
        "staff": -sum_accounts(sie.res, 0, [7690]),
        "fin": -sum_accounts(sie.res, 0, [8400, 8423]),
        "group": -sum_accounts(sie.res, 0, [8820]),
        "pfond": -sum_accounts(sie.res, 0, [8811]),
        "tax": -sum_accounts(sie.res, 0, [8910]),
    }
    result_prior_raw = {
        "net_sales": sum_account_range(sie.res, -1, 3000, 3799),
        "external": -sum_accounts(sie.res, -1, [6570]),
        "staff": -sum_accounts(sie.res, -1, [7690]),
        "fin": -sum_accounts(sie.res, -1, [8400, 8423]),
        "group": -sum_accounts(sie.res, -1, [8820]),
        "pfond": -sum_accounts(sie.res, -1, [8811]),
        "tax": -sum_accounts(sie.res, -1, [8910]),
    }

    statement_current = round_result_statement({
        "net_sales": result_current_raw["net_sales"],
        "external": result_current_raw["external"],
        "staff": result_current_raw["staff"],
        "fin": result_current_raw["fin"],
        "group": result_current_raw["group"],
        "pfond": result_current_raw["pfond"],
        "tax": result_current_raw["tax"],
    }, mode=rounding_mode)
    statement_prior = round_result_statement({
        "net_sales": result_prior_raw["net_sales"],
        "external": result_prior_raw["external"],
        "staff": result_prior_raw["staff"],
        "fin": result_prior_raw["fin"],
        "group": result_prior_raw["group"],
        "pfond": result_prior_raw["pfond"],
        "tax": result_prior_raw["tax"],
    }, mode=rounding_mode)

    c_ext = statement_current.line_items["external"]
    p_ext = statement_prior.line_items["external"]
    c_staff = statement_current.line_items["staff"]
    p_staff = statement_prior.line_items["staff"]
    c_fin = statement_current.line_items["fin"]
    p_fin = statement_prior.line_items["fin"]
    c_group = statement_current.line_items["group"]
    p_group = statement_prior.line_items["group"]
    c_pfond = statement_current.line_items["pfond"]
    p_pfond = statement_prior.line_items["pfond"]
    c_tax = statement_current.line_items["tax"]
    p_tax = statement_prior.line_items["tax"]
    c_net_sales = statement_current.line_items["net_sales"]
    p_net_sales = statement_prior.line_items["net_sales"]

    c_op = c_ext + c_staff
    p_op = p_ext + p_staff
    c_after_fin = c_op + c_fin
    p_after_fin = p_op + p_fin
    c_disp = c_group + c_pfond
    p_disp = p_group + p_pfond
    c_before_tax = c_after_fin + c_disp
    p_before_tax = p_after_fin + p_disp
    c_result = c_before_tax + c_tax
    p_result = p_before_tax + p_tax

    # Balansräkning råvärden (positiva belopp)
    asset_raw_current = {
        "shares": to_positive(get_account(sie.ub, 0, 1310)),
        "group_receivable": to_positive(get_account(sie.ub, 0, 1660)),
        "other_receivable": to_positive(get_account(sie.ub, 0, 1630)),
        "cash": to_positive(get_account(sie.ub, 0, 1930)),
    }
    asset_raw_prior = {
        "shares": to_positive(get_account(sie.ub, -1, 1310)),
        "group_receivable": to_positive(get_account(sie.ub, -1, 1660)),
        "other_receivable": to_positive(get_account(sie.ub, -1, 1630)),
        "cash": to_positive(get_account(sie.ub, -1, 1930)),
    }
    liability_raw_current = {
        "share_capital": to_positive(get_account(sie.ub, 0, 2081)),
        "retained": to_positive(get_account(sie.ub, 0, 2091)),
        "year_result_bs": to_positive(get_account(sie.ub, 0, 2099)),
        "period": to_positive(get_account(sie.ub, 0, 2110)),
        "long_group": to_positive(get_account(sie.ub, 0, 2360)),
        "other_debt": to_positive(get_account(sie.ub, 0, 2390)),
        "short_group": to_positive(get_account(sie.ub, 0, 2860)) + to_positive(get_account(sie.ub, 0, 2862)),
        "tax_debt": to_positive(get_account(sie.ub, 0, 2510)) + to_positive(get_account(sie.ub, 0, 2512)),
    }
    liability_raw_prior = {
        "share_capital": to_positive(get_account(sie.ub, -1, 2081)),
        "retained": to_positive(get_account(sie.ub, -1, 2091)),
        "year_result_bs": to_positive(get_account(sie.ub, -1, 2099)),
        "period": to_positive(get_account(sie.ub, -1, 2110)),
        "long_group": to_positive(get_account(sie.ub, -1, 2360)),
        "other_debt": to_positive(get_account(sie.ub, -1, 2390)),
        "short_group": to_positive(get_account(sie.ub, -1, 2860)) + to_positive(get_account(sie.ub, -1, 2862)),
        "tax_debt": to_positive(get_account(sie.ub, -1, 2510)) + to_positive(get_account(sie.ub, -1, 2512)),
    }

    current_balance_total_decimal_assets = sum(asset_raw_current.values(), start=ZERO)
    current_balance_total_decimal_rhs = sum(liability_raw_current.values(), start=ZERO)
    prior_balance_total_decimal_assets = sum(asset_raw_prior.values(), start=ZERO)
    prior_balance_total_decimal_rhs = sum(liability_raw_prior.values(), start=ZERO)

    shared_total_current = round_int_by_mode((current_balance_total_decimal_assets + current_balance_total_decimal_rhs) / Decimal("2"), rounding_mode)
    shared_total_prior = round_int_by_mode((prior_balance_total_decimal_assets + prior_balance_total_decimal_rhs) / Decimal("2"), rounding_mode)

    rounded_assets_current = round_balance_side(asset_raw_current, shared_total_current, rounding_mode)
    rounded_assets_prior = round_balance_side(asset_raw_prior, shared_total_prior, rounding_mode)
    rounded_liabs_current = round_balance_side(liability_raw_current, shared_total_current, rounding_mode)
    rounded_liabs_prior = round_balance_side(liability_raw_prior, shared_total_prior, rounding_mode)

    c_shares = rounded_assets_current.line_items["shares"]
    p_shares = rounded_assets_prior.line_items["shares"]
    c_receivable_group = rounded_assets_current.line_items["group_receivable"]
    p_receivable_group = rounded_assets_prior.line_items["group_receivable"]
    c_other_receivables = rounded_assets_current.line_items["other_receivable"]
    p_other_receivables = rounded_assets_prior.line_items["other_receivable"]
    c_cash = rounded_assets_current.line_items["cash"]
    p_cash = rounded_assets_prior.line_items["cash"]

    c_share_capital = rounded_liabs_current.line_items["share_capital"]
    p_share_capital = rounded_liabs_prior.line_items["share_capital"]
    c_retained = rounded_liabs_current.line_items["retained"]
    p_retained = rounded_liabs_prior.line_items["retained"]
    c_year_result_bs = rounded_liabs_current.line_items["year_result_bs"]
    p_year_result_bs = rounded_liabs_prior.line_items["year_result_bs"]
    c_period = rounded_liabs_current.line_items["period"]
    p_period = rounded_liabs_prior.line_items["period"]
    c_long_group = rounded_liabs_current.line_items["long_group"]
    p_long_group = rounded_liabs_prior.line_items["long_group"]
    c_other_debt = rounded_liabs_current.line_items["other_debt"]
    p_other_debt = rounded_liabs_prior.line_items["other_debt"]
    c_short_group = rounded_liabs_current.line_items["short_group"]
    p_short_group = rounded_liabs_prior.line_items["short_group"]
    c_tax_debt = rounded_liabs_current.line_items["tax_debt"]
    p_tax_debt = rounded_liabs_prior.line_items["tax_debt"]

    c_short_receivables = c_receivable_group + c_other_receivables
    p_short_receivables = p_receivable_group + p_other_receivables
    c_current_assets = c_short_receivables + c_cash
    p_current_assets = p_short_receivables + p_cash
    c_assets_total = c_shares + c_current_assets
    p_assets_total = p_shares + p_current_assets

    c_free_equity = c_retained + c_year_result_bs
    p_free_equity = p_retained + p_year_result_bs
    c_equity = c_share_capital + c_free_equity
    p_equity = p_share_capital + p_free_equity
    c_long_total = c_long_group + c_other_debt
    p_long_total = p_long_group + p_other_debt
    c_short_total = c_short_group + c_tax_debt
    p_short_total = p_short_group + p_tax_debt
    c_equity_and_liab_total = c_equity + c_period + c_long_total + c_short_total
    p_equity_and_liab_total = p_equity + p_period + p_long_total + p_short_total

    c_solidity = round_half_up_int(
        ((Decimal(c_equity) + Decimal(c_period) * (Decimal("1") - tax_rate)) / Decimal(c_assets_total)) * Decimal("100")
    ) if c_assets_total else 0
    p_solidity = round_half_up_int(
        ((Decimal(p_equity) + Decimal(p_period) * (Decimal("1") - tax_rate)) / Decimal(p_assets_total)) * Decimal("100")
    ) if p_assets_total else 0

    values_raw: Dict[str, Any] = {
        "current_year": current_year,
        "prior_year": prior_year,
        "company_name": manual.get("company_name", sie.company_name),
        "org_number": manual.get("org_number", sie.org_number),
        "report_start": manual.get("report_start", sie.current_start),
        "report_end": manual.get("report_end", sie.current_end),
        "prior_start": manual.get("prior_start", sie.prior_start),
        "prior_end": manual.get("prior_end", sie.prior_end),
        "adoption_date": manual.get("adoption_date", ""),
        "board_city": manual.get("board_city", ""),
        "business_description": manual.get("business_description", ""),
        "registered_seat": manual.get("registered_seat", ""),
        "avg_employees": str(manual.get("avg_employees", "")),
        "signatory_for_faststallelse": manual.get("faststallelse_signatory", ""),
        "title_for_faststallelse": manual.get("faststallelse_title", "Styrelseledamot"),
        "board_members": deepcopy(manual.get("board_members", [])),
        "historic": deepcopy(older),
        "rounding_mode": rounding_mode,
        "rounding": {
            "assets_current_target": shared_total_current,
            "assets_prior_target": shared_total_prior,
            "assets_current_adjustment_steps": rounded_assets_current.rounding_diff_from_individual,
            "assets_prior_adjustment_steps": rounded_assets_prior.rounding_diff_from_individual,
            "liabs_current_adjustment_steps": rounded_liabs_current.rounding_diff_from_individual,
            "liabs_prior_adjustment_steps": rounded_liabs_prior.rounding_diff_from_individual,
            "result_current_adjustment_steps": statement_current.rounding_diff_from_individual,
            "result_prior_adjustment_steps": statement_prior.rounding_diff_from_individual,
            "raw_assets_total_current": str(current_balance_total_decimal_assets),
            "raw_liabs_total_current": str(current_balance_total_decimal_rhs),
            "raw_assets_total_prior": str(prior_balance_total_decimal_assets),
            "raw_liabs_total_prior": str(prior_balance_total_decimal_rhs),
        },
        "result": {
            "net_sales_current": c_net_sales,
            "net_sales_prior": p_net_sales,
            "external_current": c_ext,
            "external_prior": p_ext,
            "staff_current": c_staff,
            "staff_prior": p_staff,
            "op_current": c_op,
            "op_prior": p_op,
            "fin_current": c_fin,
            "fin_prior": p_fin,
            "after_fin_current": c_after_fin,
            "after_fin_prior": p_after_fin,
            "group_current": c_group,
            "group_prior": p_group,
            "pfond_current": c_pfond,
            "pfond_prior": p_pfond,
            "disp_current": c_disp,
            "disp_prior": p_disp,
            "before_tax_current": c_before_tax,
            "before_tax_prior": p_before_tax,
            "tax_current": c_tax,
            "tax_prior": p_tax,
            "year_current": c_result,
            "year_prior": p_result,
        },
        "balance": {
            "shares_current": c_shares,
            "shares_prior": p_shares,
            "group_receivable_current": c_receivable_group,
            "group_receivable_prior": p_receivable_group,
            "other_receivable_current": c_other_receivables,
            "other_receivable_prior": p_other_receivables,
            "cash_current": c_cash,
            "cash_prior": p_cash,
            "short_receivables_current": c_short_receivables,
            "short_receivables_prior": p_short_receivables,
            "current_assets_current": c_current_assets,
            "current_assets_prior": p_current_assets,
            "assets_total_current": c_assets_total,
            "assets_total_prior": p_assets_total,
            "share_capital_current": c_share_capital,
            "share_capital_prior": p_share_capital,
            "retained_current": c_retained,
            "retained_prior": p_retained,
            "year_result_bs_current": c_year_result_bs,
            "year_result_bs_prior": p_year_result_bs,
            "free_equity_current": c_free_equity,
            "free_equity_prior": p_free_equity,
            "equity_current": c_equity,
            "equity_prior": p_equity,
            "period_current": c_period,
            "period_prior": p_period,
            "long_group_current": c_long_group,
            "long_group_prior": p_long_group,
            "other_debt_current": c_other_debt,
            "other_debt_prior": p_other_debt,
            "long_total_current": c_long_total,
            "long_total_prior": p_long_total,
            "short_group_current": c_short_group,
            "short_group_prior": p_short_group,
            "tax_debt_current": c_tax_debt,
            "tax_debt_prior": p_tax_debt,
            "short_total_current": c_short_total,
            "short_total_prior": p_short_total,
            "equity_and_liab_total_current": c_equity_and_liab_total,
            "equity_and_liab_total_prior": p_equity_and_liab_total,
        },
        "metrics": {
            "solidity_current": c_solidity,
            "solidity_prior": p_solidity,
        },
    }

    apply_overrides(values_raw, manual.get("overrides", {}))

    year_minus_2 = str(report_year - 2)
    year_minus_3 = str(report_year - 3)
    older_a = older.get(year_minus_2, {})
    older_b = older.get(year_minus_3, {})

    tables = {
        "table0": [
            ["", values_raw["report_start"], values_raw["prior_start"], older_a.get("start", f"{year_minus_2}-01-01"), older_b.get("start", f"{year_minus_3}-01-01")],
            ["", f"- {values_raw['report_end']}", f"- {values_raw['prior_end']}", f"- {older_a.get('end', f'{year_minus_2}-12-31')}", f"- {older_b.get('end', f'{year_minus_3}-12-31')}"] ,
            ["Nettoomsättning", format_kr(values_raw["result"]["net_sales_current"]), format_kr(values_raw["result"]["net_sales_prior"]), format_kr(older_a.get("nettoomsattning", 0)), format_kr(older_b.get("nettoomsattning", 0))],
            ["Resultat efter finansiella poster", format_kr(values_raw["result"]["after_fin_current"]), format_kr(values_raw["result"]["after_fin_prior"]), format_kr(older_a.get("resultat_efter_finansiella_poster", 0)), format_kr(older_b.get("resultat_efter_finansiella_poster", 0))],
            ["Soliditet (%)", str(values_raw["metrics"]["solidity_current"]), str(values_raw["metrics"]["solidity_prior"]), str(older_a.get("soliditet", "")), str(older_b.get("soliditet", ""))],
            ["Vid årets ingång", format_kr(values_raw["balance"]["share_capital_prior"]), format_kr(values_raw["balance"]["retained_prior"]), format_kr(values_raw["balance"]["year_result_bs_prior"]), format_kr(values_raw["balance"]["equity_prior"])],
            ["Balanseras i ny räkning", "", format_kr(values_raw["balance"]["year_result_bs_prior"]), format_kr(-values_raw["balance"]["year_result_bs_prior"]), format_kr(0)],
            ["Årets resultat", "", "", format_kr(values_raw["balance"]["year_result_bs_current"]), format_kr(values_raw["balance"]["year_result_bs_current"])],
            ["Vid årets utgång", format_kr(values_raw["balance"]["share_capital_current"]), format_kr(values_raw["balance"]["retained_current"]), format_kr(values_raw["balance"]["year_result_bs_current"]), format_kr(values_raw["balance"]["equity_current"])],
            ["Balanserat resultat", format_kr(values_raw["balance"]["retained_current"]), format_kr(values_raw["balance"]["retained_current"]), format_kr(values_raw["balance"]["retained_current"])],
            ["Årets resultat", format_kr(values_raw["balance"]["year_result_bs_current"]), format_kr(values_raw["balance"]["year_result_bs_current"]), format_kr(values_raw["balance"]["year_result_bs_current"])],
            ["Summa", format_kr(values_raw["balance"]["free_equity_current"]), format_kr(values_raw["balance"]["free_equity_current"]), format_kr(values_raw["balance"]["free_equity_current"])],
            ["Utdelas till aktieägare", format_kr(manual.get("utdelning", 0)), format_kr(manual.get("utdelning", 0)), format_kr(manual.get("utdelning", 0))],
            ["Balanseras i ny räkning", format_kr(values_raw["balance"]["free_equity_current"] - int(manual.get("utdelning", 0))), format_kr(values_raw["balance"]["free_equity_current"] - int(manual.get("utdelning", 0))), format_kr(values_raw["balance"]["free_equity_current"] - int(manual.get("utdelning", 0)))],
            ["Summa", format_kr(values_raw["balance"]["free_equity_current"]), format_kr(values_raw["balance"]["free_equity_current"]), format_kr(values_raw["balance"]["free_equity_current"])],
        ],
        "table1": [
            ["Övriga externa kostnader", format_kr(values_raw["result"]["external_current"]), format_kr(values_raw["result"]["external_prior"])],
            ["Personalkostnader", format_kr(values_raw["result"]["staff_current"]), format_kr(values_raw["result"]["staff_prior"])],
            ["Summa rörelsekostnader", format_kr(values_raw["result"]["op_current"]), format_kr(values_raw["result"]["op_prior"])],
            ["Rörelseresultat", format_kr(values_raw["result"]["op_current"]), format_kr(values_raw["result"]["op_prior"])],
            ["Räntekostnader och liknande resultatposter", format_kr(values_raw["result"]["fin_current"]), format_kr(values_raw["result"]["fin_prior"])],
            ["Summa finansiella poster", format_kr(values_raw["result"]["fin_current"]), format_kr(values_raw["result"]["fin_prior"])],
            ["Resultat efter finansiella poster", format_kr(values_raw["result"]["after_fin_current"]), format_kr(values_raw["result"]["after_fin_prior"])],
            ["Erhållna koncernbidrag", format_kr(values_raw["result"]["group_current"]), format_kr(values_raw["result"]["group_prior"])],
            ["Förändring av periodiseringsfonder", format_kr(values_raw["result"]["pfond_current"]), format_kr(values_raw["result"]["pfond_prior"])],
            ["Summa bokslutsdispositioner", format_kr(values_raw["result"]["disp_current"]), format_kr(values_raw["result"]["disp_prior"])],
            ["Resultat före skatt", format_kr(values_raw["result"]["before_tax_current"]), format_kr(values_raw["result"]["before_tax_prior"])],
            ["Skatt på årets resultat", format_kr(values_raw["result"]["tax_current"]), format_kr(values_raw["result"]["tax_prior"])],
            ["Årets resultat", format_kr(values_raw["result"]["year_current"]), format_kr(values_raw["result"]["year_prior"])],
        ],
        "table2": [
            ["Andelar i koncernföretag", format_kr(values_raw["balance"]["shares_current"]), format_kr(values_raw["balance"]["shares_prior"])],
            ["Summa finansiella anläggningstillgångar", format_kr(values_raw["balance"]["shares_current"]), format_kr(values_raw["balance"]["shares_prior"])],
            ["Summa anläggningstillgångar", format_kr(values_raw["balance"]["shares_current"]), format_kr(values_raw["balance"]["shares_prior"])],
            ["Fordringar hos koncernföretag", format_kr(values_raw["balance"]["group_receivable_current"]), format_kr(values_raw["balance"]["group_receivable_prior"])],
            ["Övriga fordringar", format_kr(values_raw["balance"]["other_receivable_current"]), format_kr(values_raw["balance"]["other_receivable_prior"])],
            ["Summa kortfristiga fordringar", format_kr(values_raw["balance"]["short_receivables_current"]), format_kr(values_raw["balance"]["short_receivables_prior"])],
            ["Kassa och bank", format_kr(values_raw["balance"]["cash_current"]), format_kr(values_raw["balance"]["cash_prior"])],
            ["Summa kassa och bank", format_kr(values_raw["balance"]["cash_current"]), format_kr(values_raw["balance"]["cash_prior"])],
            ["Summa omsättningstillgångar", format_kr(values_raw["balance"]["current_assets_current"]), format_kr(values_raw["balance"]["current_assets_prior"])],
            ["Summa tillgångar", format_kr(values_raw["balance"]["assets_total_current"]), format_kr(values_raw["balance"]["assets_total_prior"])],
        ],
        "table3": [
            ["Aktiekapital", format_kr(values_raw["balance"]["share_capital_current"]), format_kr(values_raw["balance"]["share_capital_prior"])],
            ["Summa bundet eget kapital", format_kr(values_raw["balance"]["share_capital_current"]), format_kr(values_raw["balance"]["share_capital_prior"])],
            ["Balanserat resultat", format_kr(values_raw["balance"]["retained_current"]), format_kr(values_raw["balance"]["retained_prior"])],
            ["Årets resultat", format_kr(values_raw["balance"]["year_result_bs_current"]), format_kr(values_raw["balance"]["year_result_bs_prior"])],
            ["Summa fritt eget kapital", format_kr(values_raw["balance"]["free_equity_current"]), format_kr(values_raw["balance"]["free_equity_prior"])],
            ["Summa eget kapital", format_kr(values_raw["balance"]["equity_current"]), format_kr(values_raw["balance"]["equity_prior"])],
            ["Periodiseringsfonder", format_kr(values_raw["balance"]["period_current"]), format_kr(values_raw["balance"]["period_prior"])],
            ["Summa obeskattade reserver", format_kr(values_raw["balance"]["period_current"]), format_kr(values_raw["balance"]["period_prior"])],
            ["Skulder till koncernföretag_lång", format_kr(values_raw["balance"]["long_group_current"]), format_kr(values_raw["balance"]["long_group_prior"])],
            ["Övriga skulder", format_kr(values_raw["balance"]["other_debt_current"]), format_kr(values_raw["balance"]["other_debt_prior"])],
            ["Summa långfristiga skulder", format_kr(values_raw["balance"]["long_total_current"]), format_kr(values_raw["balance"]["long_total_prior"])],
            ["Skulder till koncernföretag_kort", format_kr(values_raw["balance"]["short_group_current"]), format_kr(values_raw["balance"]["short_group_prior"])],
            ["Skatteskulder", format_kr(values_raw["balance"]["tax_debt_current"]), format_kr(values_raw["balance"]["tax_debt_prior"])],
            ["Summa kortfristiga skulder", format_kr(values_raw["balance"]["short_total_current"]), format_kr(values_raw["balance"]["short_total_prior"])],
            ["Summa eget kapital och skulder", format_kr(values_raw["balance"]["equity_and_liab_total_current"]), format_kr(values_raw["balance"]["equity_and_liab_total_prior"])],
        ],
    }

    paragraphs = {
        "company_name_cover": values_raw["company_name"],
        "org_cover": values_raw["org_number"],
        "cover_period": f"Årsredovisning för räkenskapsåret {values_raw['report_start']} - {values_raw['report_end']}",
        "faststallelse": (
            f"Undertecknad {values_raw['title_for_faststallelse'].lower()} i {values_raw['company_name']} "
            f"intygar att resultaträkningen och balansräkningen har fastställts på årsstämman den "
            f"{values_raw['adoption_date']}. Årsstämman beslutade att godkänna styrelsens förslag till hur vinsten ska disponeras."
        ),
        "faststallelse_sign": f"{values_raw['signatory_for_faststallelse']}, {values_raw['title_for_faststallelse']}",
        "faststallelse_city_date": f"{values_raw['board_city']} {values_raw['adoption_date']}",
        "business_text": (
            f"{values_raw['business_description']} Företaget har sitt säte i {values_raw['registered_seat']}."
            if values_raw["business_description"] else ""
        ),
        "note2_start": values_raw["report_start"],
        "note2_end": f"- {values_raw['report_end']}",
        "note2_avg": f"Medelantal anställda under året\t{values_raw['avg_employees']}",
        "note3_date": values_raw["report_end"],
        "note3_ing": f"Ingående anskaffningsvärden\t{format_kr(values_raw['balance']['shares_current'])}",
        "note3_utg": f"Utgående anskaffningsvärden\t{format_kr(values_raw['balance']['shares_current'])}",
        "note3_book": f"Redovisat värde\t{format_kr(values_raw['balance']['shares_current'])}",
        "signature_heading": (
            f"Årsredovisning för {values_raw['company_name']}, {values_raw['org_number']} "
            f"Avseende räkenskapsåret {values_raw['report_start']} - {values_raw['report_end']}"
        ),
    }

    return BuiltValues(paragraphs=paragraphs, tables=tables, raw=values_raw)





def paragraph_text(paragraph) -> str:
    return paragraph.text.replace("\xa0", " ").strip()


def find_first_paragraph(doc: Document, predicate) -> Optional[Paragraph]:
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_placeholder_text_everywhere(doc: Document, context: Dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, context)


def replace_placeholders_in_paragraph(paragraph, context: Dict[str, str]) -> None:
    text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if not text:
        return
    new_text = text
    for key, value in context.items():
        new_text = new_text.replace("{{" + key + "}}", value or "")
    if new_text != text:
        replace_runs(paragraph, new_text)


RAW_TEMPLATE_PARAGRAPH_MAP = {
    1: "{{company_name_cover}}",
    2: "{{org_number_cover}}",
    5: "Årsredovisning för räkenskapsåret {{report_start}} - {{report_end}}",
    17: "{{faststallelse_text}}",
    29: "{{faststallelse_sign_line}}",
    30: "{{faststallelse_city_date}}",
    41: "{{business_text}}",
    52: "{{note1_main}}",
    64: "{{note2_start}}",
    65: "{{note2_end_line}}",
    66: "Medelantal anställda under året\t{{avg_employees}}",
    69: "{{note3_date}}",
    71: "Ingående anskaffningsvärden\t{{note3_opening}}",
    72: "Utgående anskaffningsvärden\t{{note3_closing}}",
    73: "Redovisat värde\t{{note3_book}}",
    76: "Årsredovisning för {{company_name}}, {{org_number}} Avseende räkenskapsåret {{report_start}} - {{report_end}}",
    80: "{{signature_city}}",
    86: "{{BLOCK_SIGNATURES}}",
}


def prepare_managed_template(doc: Document) -> Document:
    if any("{{company_name_cover}}" in paragraph.text for paragraph in doc.paragraphs):
        return doc

    for idx, text in RAW_TEMPLATE_PARAGRAPH_MAP.items():
        if idx < len(doc.paragraphs):
            replace_runs(doc.paragraphs[idx], text)

    # Capture original anchor paragraphs before inserting new ones so indexes do not drift.
    business_para = doc.paragraphs[41] if len(doc.paragraphs) > 41 else None
    note1_para = doc.paragraphs[52] if len(doc.paragraphs) > 52 else None
    note3_para = doc.paragraphs[73] if len(doc.paragraphs) > 73 else None

    # Remove old signature date/name paragraphs after the anchor to avoid duplicate blocks
    for idx in sorted([87, 88, 89], reverse=True):
        if idx < len(doc.paragraphs):
            remove_paragraph(doc.paragraphs[idx])

    # Insert anchors in descending document order.
    if note3_para is not None and not any("{{BLOCK_NOTES_EXTRA}}" in p.text for p in doc.paragraphs):
        insert_paragraph_after(note3_para, "{{BLOCK_NOTES_EXTRA}}", "Body Text")

    if note1_para is not None and not any("{{BLOCK_NOTE1_EXTRA}}" in p.text for p in doc.paragraphs):
        insert_paragraph_after(note1_para, "{{BLOCK_NOTE1_EXTRA}}", "Body Text")

    if business_para is not None and not any("{{BLOCK_FORVALTNING_EXTRA}}" in p.text for p in doc.paragraphs):
        insert_paragraph_after(business_para, "{{BLOCK_FORVALTNING_EXTRA}}", "Body Text")

    return doc


def build_placeholder_context(built: BuiltValues) -> Dict[str, str]:
    raw = built.raw
    note1_main = "Årsredovisningen är upprättad i enlighet med årsredovisningslagen och Bokföringsnämndens allmänna råd (BFNAR 2016:10) om årsredovisning i mindre företag."
    return {
        "company_name_cover": raw["company_name"],
        "org_number_cover": raw["org_number"],
        "company_name": raw["company_name"],
        "org_number": raw["org_number"],
        "report_start": raw["report_start"],
        "report_end": raw["report_end"],
        "faststallelse_text": (
            f"Undertecknad {raw['title_for_faststallelse'].lower()} i {raw['company_name']} intygar att resultaträkningen och balansräkningen har fastställts på årsstämman den {raw['adoption_date']}. Årsstämman beslutade att godkänna styrelsens förslag till hur vinsten ska disponeras."
        ),
        "faststallelse_sign_line": f"{raw['signatory_for_faststallelse']}, {raw['title_for_faststallelse']}",
        "faststallelse_city_date": f"{raw['board_city']} {raw['adoption_date']}".strip(),
        "business_text": (f"{raw['business_description']} Företaget har sitt säte i {raw['registered_seat']}." if raw["business_description"] else ""),
        "note1_main": note1_main,
        "note2_start": raw["report_start"],
        "note2_end_line": f"- {raw['report_end']}",
        "avg_employees": raw["avg_employees"],
        "note3_date": raw["report_end"],
        "note3_opening": format_kr(raw['balance']['shares_current']),
        "note3_closing": format_kr(raw['balance']['shares_current']),
        "note3_book": format_kr(raw['balance']['shares_current']),
        "signature_city": raw["board_city"],
    }


def build_forvaltnings_sections(manual: Dict[str, Any]) -> List[Dict[str, str]]:
    blocks: List[Dict[str, str]] = []
    if str(manual.get("significant_events_during_year", "")).strip():
        blocks.append({"style": "Heading 3", "text": "Väsentliga händelser under räkenskapsåret"})
        blocks.append({"style": "Body Text", "text": str(manual.get("significant_events_during_year", "")).strip()})
    if manual.get("has_own_shares"):
        blocks.append({"style": "Heading 3", "text": "Egna aktier"})
        blocks.append({"style": "Body Text", "text": str(manual.get("own_shares_text", "")).strip()})
    if str(manual.get("net_sales_variation_comment", "")).strip():
        blocks.append({"style": "Heading 3", "text": "Kommentar till flerårsöversikten"})
        blocks.append({"style": "Body Text", "text": str(manual.get("net_sales_variation_comment", "")).strip()})
    return blocks


def build_note1_extra_blocks(manual: Dict[str, Any]) -> List[Dict[str, str]]:
    blocks: List[Dict[str, str]] = []
    if not manual.get("k2_previous_year", True):
        blocks.append({
            "style": "Body Text",
            "text": "Årsredovisningen upprättas för första gången i enlighet med Bokföringsnämndens allmänna råd (BFNAR 2016:10) om årsredovisning i mindre företag, vilket kan innebära en bristande jämförbarhet mellan räkenskapsåret och det närmast föregående räkenskapsåret."
        })
    return blocks


def build_extra_notes(sie: SieData, built: BuiltValues, manual: Dict[str, Any]) -> List[Dict[str, str]]:
    blocks: List[Dict[str, str]] = []
    note_no = 4
    if str(manual.get("significant_events_after_year_end", "")).strip():
        blocks.append({"style": "Heading 2", "text": f"Not {note_no} - Väsentliga händelser efter räkenskapsårets slut"})
        blocks.append({"style": "Body Text", "text": str(manual.get("significant_events_after_year_end", "")).strip()})
        note_no += 1

    group_interest_current = to_positive(get_account(sie.res, 0, 8423))
    group_interest_prior = to_positive(get_account(sie.res, -1, 8423))
    if group_interest_current or group_interest_prior:
        blocks.append({"style": "Heading 2", "text": f"Not {note_no} - Räntekostnader till koncernföretag"})
        blocks.append({"style": "Body Text", "text": f"Av årets räntekostnader och liknande resultatposter avser {format_kr(round_int_by_mode(group_interest_current, str(manual.get('rounding_mode', 'truncate'))))} kr ({format_kr(round_int_by_mode(group_interest_prior, str(manual.get('rounding_mode', 'truncate'))))} kr) skulder till koncernföretag."})
    return blocks


def expand_anchor_paragraph(doc: Document, anchor_text: str, blocks: List[Dict[str, str]]) -> None:
    anchor = find_first_paragraph(doc, lambda p: anchor_text in p.text)
    if anchor is None:
        return
    current = anchor
    for block in blocks:
        current = insert_paragraph_after(current, block.get("text", ""), block.get("style"))
    remove_paragraph(anchor)


def populate_signature_block(doc: Document, built: BuiltValues) -> None:
    anchor = find_first_paragraph(doc, lambda p: "{{BLOCK_SIGNATURES}}" in p.text)
    if anchor is None:
        return
    current = anchor
    members = built.raw.get("board_members", []) or []
    for idx, member in enumerate(members):
        name_line = " ".join(part for part in [str(member.get("name", "")).strip(), str(member.get("title", "")).strip()] if part).strip()
        date_line = str(member.get("date", built.raw.get("adoption_date", ""))).strip()
        current = insert_paragraph_after(current, "", "Body Text")
        current = insert_paragraph_after(current, "", "Body Text")
        current = insert_paragraph_after(current, name_line, "Body Text")
        current = insert_paragraph_after(current, date_line, "Body Text")
    remove_paragraph(anchor)


def fill_tables(doc: Document, built: BuiltValues) -> None:
    t0 = doc.tables[0]
    table0 = built.tables["table0"]
    for col in range(1, 5):
        set_cell_text(t0.rows[1].cells[col], table0[0][col])
        set_cell_text(t0.rows[2].cells[col], table0[1][col])
    for row_idx, source_idx in [(3, 2), (4, 3), (5, 4), (8, 5), (9, 6), (10, 7), (11, 8)]:
        row = t0.rows[row_idx]
        data = table0[source_idx]
        for col, value in enumerate(data[1:], start=1):
            set_cell_text(row.cells[col], value)
    for row_idx, source_idx in [(14, 9), (15, 10), (16, 11), (18, 12), (19, 13), (20, 14)]:
        row = t0.rows[row_idx]
        data = table0[source_idx]
        for col in (2, 3, 4):
            set_cell_text(row.cells[col], data[col - 1])

    t1 = doc.tables[1]
    set_cell_text(t1.rows[1].cells[2], built.raw["report_start"])
    set_cell_text(t1.rows[1].cells[3], built.raw["prior_start"])
    set_cell_text(t1.rows[2].cells[2], f"- {built.raw['report_end']}")
    set_cell_text(t1.rows[2].cells[3], f"- {built.raw['prior_end']}")
    row_map_1 = {4: 0, 5: 1, 6: 2, 7: 3, 9: 4, 10: 5, 11: 6, 13: 7, 14: 8, 15: 9, 16: 10, 18: 11, 19: 12}
    for row_idx, source_idx in row_map_1.items():
        data = built.tables["table1"][source_idx]
        set_cell_text(t1.rows[row_idx].cells[2], data[1])
        set_cell_text(t1.rows[row_idx].cells[3], data[2])

    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[2], built.raw["report_end"])
    set_cell_text(t2.rows[1].cells[3], built.raw["prior_end"])
    row_map_2 = {5: 0, 6: 1, 7: 2, 10: 3, 11: 4, 12: 5, 14: 6, 15: 7, 16: 8, 17: 9}
    for row_idx, source_idx in row_map_2.items():
        data = built.tables["table2"][source_idx]
        set_cell_text(t2.rows[row_idx].cells[2], data[1])
        set_cell_text(t2.rows[row_idx].cells[3], data[2])

    t3 = doc.tables[3]
    set_cell_text(t3.rows[1].cells[2], built.raw["report_end"])
    set_cell_text(t3.rows[1].cells[3], built.raw["prior_end"])
    row_map_3 = {5: 0, 6: 1, 8: 2, 9: 3, 10: 4, 11: 5, 13: 6, 14: 7, 16: 8, 17: 9, 18: 10, 20: 11, 21: 12, 22: 13, 23: 14}
    for row_idx, source_idx in row_map_3.items():
        data = built.tables["table3"][source_idx]
        set_cell_text(t3.rows[row_idx].cells[2], data[1])
        set_cell_text(t3.rows[row_idx].cells[3], data[2])




# --- Layout-focused v7 helpers ---

def clear_paragraph(paragraph) -> None:
    if paragraph._p.getparent() is None:
        return
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def apply_run_font(run, font_name: str = "Times New Roman", size_pt: float = 11.0, bold: Optional[bool] = None, italic: Optional[bool] = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) if run._element.rPr is not None else None
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(paragraph, text: str = '', *, bold: bool = False, italic: bool = False, size: float = 11.0, font_name: str = 'Times New Roman'):
    run = paragraph.add_run(text)
    apply_run_font(run, font_name=font_name, size_pt=size, bold=bold, italic=italic)
    return run


def set_paragraph_base(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.0, keep_with_next=False, keep_together=False, left_indent_cm=None, first_line_indent_cm=None):
    fmt = paragraph.paragraph_format
    paragraph.alignment = align
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)


def add_heading_clean(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_base(p, space_before=0, space_after=10, keep_with_next=True)
        add_text(p, text, bold=True, size=18)
    elif level == 2:
        set_paragraph_base(p, space_before=6, space_after=6, keep_with_next=True)
        add_text(p, text, bold=True, size=14)
    elif level == 3:
        set_paragraph_base(p, space_before=4, space_after=2, keep_with_next=True)
        add_text(p, text, bold=True, size=12)
    else:
        set_paragraph_base(p, space_before=2, space_after=2, keep_with_next=True)
        add_text(p, text, bold=True, size=11)
    return p


def add_body_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=5, size=11.0):
    p = doc.add_paragraph()
    set_paragraph_base(p, align=align, space_before=space_before, space_after=space_after)
    add_text(p, text, size=size)
    return p


def set_document_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    styles = doc.styles
    for style_name in ['Normal', 'Body Text']:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.font.size = Pt(11)
    for style_name, size in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12), ('Heading 4', 11)]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.font.size = Pt(size)
            style.font.bold = True


def add_field_run(paragraph, instr: str):
    run = paragraph.add_run()
    apply_run_font(run, size_pt=10.0)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    default_text = OxmlElement('w:t')
    default_text.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(default_text)
    run._r.append(fld_end)
    return run


def set_top_border(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    top = pBdr.find(qn('w:top'))
    if top is None:
        top = OxmlElement('w:top')
        pBdr.append(top)
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), '666666')


def add_footer_clean(doc: Document, company_name: str, org_number: str) -> None:
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(p)
        set_paragraph_base(p, space_before=0, space_after=0, line_spacing=1.0)
        set_top_border(p)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
        add_text(p, f"{company_name} {org_number}", size=10)
        add_text(p, '	Sida ', size=10)
        add_field_run(p, 'PAGE')
        add_text(p, ' av ', size=10)
        add_field_run(p, 'NUMPAGES')


def set_cell_margins(cell, top=60, start=90, bottom=60, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tcMar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_borders(table, *, top=True, bottom=True, inside_h=False, inside_v=False, left=False, right=False):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    options = {'top': top, 'bottom': bottom, 'left': left, 'right': right, 'insideH': inside_h, 'insideV': inside_v}
    for key, enabled in options.items():
        node = borders.find(qn(f'w:{key}'))
        if node is None:
            node = OxmlElement(f'w:{key}')
            borders.append(node)
        if enabled:
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), '6')
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), '808080')
        else:
            node.set(qn('w:val'), 'nil')


def set_cell_text_clean(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11.0):
    cell.text = ''
    p = cell.paragraphs[0]
    set_paragraph_base(p, align=align, space_before=0, space_after=0, line_spacing=1.0)
    add_text(p, text or '', bold=bold, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def configure_table(table, widths_cm, *, alignment=WD_TABLE_ALIGNMENT.CENTER, font_size=10.5, first_row_bold=False, top_bottom_borders=True, full_grid=False):
    table.alignment = alignment
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells[:len(widths_cm)]):
            cell.width = Cm(widths_cm[idx])
            for p in cell.paragraphs:
                for r in p.runs:
                    apply_run_font(r, size_pt=font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table, top=top_bottom_borders, bottom=top_bottom_borders, inside_h=full_grid, inside_v=False, left=False, right=False)
    if first_row_bold and table.rows:
        for c in table.rows[0].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True


def add_spacer(doc: Document, height_pt: float = 6.0):
    p = doc.add_paragraph()
    set_paragraph_base(p, space_before=0, space_after=0, line_spacing=1.0)
    run = add_text(p, '', size=1)
    p.paragraph_format.space_after = Pt(height_pt)
    return p


def add_cover_page(doc: Document, built: BuiltValues) -> None:
    raw = built.raw
    for _ in range(4):
        add_spacer(doc, 10)
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)
    add_text(p, raw['company_name'], bold=True, size=18)
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_text(p, raw['org_number'], size=12)

    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=8)
    add_text(p, 'Årsredovisning', bold=True, size=20)
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_text(p, 'för räkenskapsåret', size=12)
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_text(p, f"{raw['report_start']} - {raw['report_end']}", size=14)

    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=0)
    add_text(p, 'Styrelsen upprättar följande årsredovisning.', size=11)
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_text(p, 'Samtliga belopp är angivna i hela kronor.', size=11)


def add_faststallelse_page(doc: Document, built: BuiltValues) -> None:
    raw = built.raw
    add_heading_clean(doc, 'Fastställelseintyg', 1)
    add_body_paragraph(doc, (
        f"Undertecknad {raw['title_for_faststallelse'].lower()} i {raw['company_name']} intygar att resultaträkningen och balansräkningen har fastställts på årsstämman den {raw['adoption_date']}. Årsstämman beslutade att godkänna styrelsens förslag till hur vinsten ska disponeras."
    ), space_after=8)
    add_body_paragraph(doc, 'Jag intygar också att innehållet i årsredovisningen stämmer överens med originalet.', space_after=20)
    add_body_paragraph(doc, f"{raw['board_city']} {raw['adoption_date']}", space_after=18)
    add_body_paragraph(doc, f"{raw['signatory_for_faststallelse']}, {raw['title_for_faststallelse']}", space_after=0)


def add_flerarsoversikt_table(doc: Document, table0):
    table = doc.add_table(rows=5, cols=5)
    configure_table(table, [6.4, 2.6, 2.6, 2.6, 2.6], font_size=10.0, top_bottom_borders=True)
    for col in range(5):
        text = table0[0][col] if col < len(table0[0]) else ''
        set_cell_text_clean(table.cell(0, col), text, align=WD_ALIGN_PARAGRAPH.CENTER if col else WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10.0)
        text2 = table0[1][col] if col < len(table0[1]) else ''
        set_cell_text_clean(table.cell(1, col), text2, align=WD_ALIGN_PARAGRAPH.CENTER if col else WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10.0)
    for r in range(3):
        source = table0[2 + r]
        set_cell_text_clean(table.cell(2 + r, 0), source[0], align=WD_ALIGN_PARAGRAPH.LEFT, size=10.0)
        for c in range(1, 5):
            set_cell_text_clean(table.cell(2 + r, c), source[c], align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.0)
    return table


def add_equity_change_table(doc: Document, table0):
    table = doc.add_table(rows=5, cols=5)
    configure_table(table, [6.4, 2.5, 2.7, 2.5, 2.5], font_size=10.0, top_bottom_borders=True)
    headers = ['', 'Aktiekapital', 'Balanserat resultat', 'Årets resultat', 'Totalt']
    for c, text in enumerate(headers):
        set_cell_text_clean(table.cell(0, c), text, align=WD_ALIGN_PARAGRAPH.CENTER if c else WD_ALIGN_PARAGRAPH.LEFT, size=10.0, bold=False)
    for r in range(4):
        src = table0[5 + r]
        set_cell_text_clean(table.cell(1 + r, 0), src[0], align=WD_ALIGN_PARAGRAPH.LEFT, size=10.0)
        for c in range(1, 5):
            set_cell_text_clean(table.cell(1 + r, c), src[c], align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.0)
    return table


def add_result_disposition_table(doc: Document, table0):
    add_body_paragraph(doc, 'Styrelsen föreslår att till förfogande stående medel', space_after=6)
    table = doc.add_table(rows=6, cols=2)
    configure_table(table, [9.5, 4.2], font_size=11.0, top_bottom_borders=False)
    rows = [table0[9], table0[10], table0[11], ['', ''], table0[12], table0[13]]
    labels = ['Balanserat resultat', 'Årets resultat', 'Summa', 'Disponeras enligt följande', 'Utdelas till aktieägare', 'Balanseras i ny räkning']
    values = [table0[9][1], table0[10][1], table0[11][1], '', table0[12][1], table0[13][1]]
    for i in range(6):
        set_cell_text_clean(table.cell(i, 0), labels[i], align=WD_ALIGN_PARAGRAPH.LEFT, size=11.0, bold=(labels[i] == 'Summa'))
        set_cell_text_clean(table.cell(i, 1), values[i], align=WD_ALIGN_PARAGRAPH.RIGHT, size=11.0, bold=(labels[i] == 'Summa'))
    add_spacer(doc, 2)
    p = doc.add_paragraph()
    set_paragraph_base(p, space_before=0, space_after=0)
    add_text(p, 'Summa', bold=True)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
    add_text(p, '	' + table0[14][1], bold=True)
    return table


def add_two_year_statement_table(doc: Document, heading: str, current_label: str, prior_label: str, rows: list, section_breaks: list, *, note_label: Optional[str] = None, font_size=10.5):
    add_heading_clean(doc, heading, 1)
    cols = 4 if note_label else 3
    widths = [7.8, 1.0, 3.0, 3.0] if note_label else [8.8, 3.2, 3.2]

    expanded_rows = []
    breaks = sorted(section_breaks, key=lambda x: x[0])
    break_idx = 0
    for idx, row in enumerate(rows):
        while break_idx < len(breaks) and breaks[break_idx][0] == idx:
            expanded_rows.append(('group', breaks[break_idx][1], '', ''))
            break_idx += 1
        expanded_rows.append(('data', row[0], row[1], row[2]))
    while break_idx < len(breaks):
        expanded_rows.append(('group', breaks[break_idx][1], '', ''))
        break_idx += 1

    table = doc.add_table(rows=1 + len(expanded_rows), cols=cols)
    configure_table(table, widths, font_size=font_size, top_bottom_borders=True)
    header = table.rows[0].cells
    set_cell_text_clean(header[0], '', size=font_size)
    offset = 1
    if note_label:
        set_cell_text_clean(header[1], 'Not', align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
        offset = 2
    set_cell_text_clean(header[offset], current_label, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
    set_cell_text_clean(header[offset + 1], prior_label, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)

    for idx, row in enumerate(expanded_rows, start=1):
        row_type, label, current_value, prior_value = row
        if row_type == 'group':
            set_cell_text_clean(table.cell(idx, 0), label, bold=True, size=font_size)
            if note_label:
                set_cell_text_clean(table.cell(idx, 1), '', align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
                set_cell_text_clean(table.cell(idx, 2), '', align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size)
                set_cell_text_clean(table.cell(idx, 3), '', align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size)
            else:
                set_cell_text_clean(table.cell(idx, 1), '', align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size)
                set_cell_text_clean(table.cell(idx, 2), '', align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size)
            continue

        bold_value = label.startswith('Summa') or label in ('Rörelseresultat', 'Resultat efter finansiella poster', 'Resultat före skatt', 'Årets resultat', 'Summa tillgångar', 'Summa eget kapital och skulder')
        set_cell_text_clean(table.cell(idx, 0), label, size=font_size, bold=bold_value)
        if note_label:
            note = ''
            if heading == 'Resultaträkning' and label == 'Övriga externa kostnader':
                note = '1'
            elif heading == 'Resultaträkning' and label == 'Personalkostnader':
                note = '2'
            elif heading == 'Balansräkning' and label == 'Andelar i koncernföretag':
                note = '3'
            set_cell_text_clean(table.cell(idx, 1), note, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
            set_cell_text_clean(table.cell(idx, 2), current_value, align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size, bold=bold_value)
            set_cell_text_clean(table.cell(idx, 3), prior_value, align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size, bold=bold_value)
        else:
            set_cell_text_clean(table.cell(idx, 1), current_value, align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size, bold=bold_value)
            set_cell_text_clean(table.cell(idx, 2), prior_value, align=WD_ALIGN_PARAGRAPH.RIGHT, size=font_size, bold=bold_value)
    return table


def add_notes_page(doc: Document, built: BuiltValues, manual: Dict[str, Any], sie: SieData):
    raw = built.raw
    add_heading_clean(doc, 'Noter', 1)

    add_heading_clean(doc, 'Not 1 - Redovisningsprinciper', 2)
    add_body_paragraph(doc, 'Årsredovisningen är upprättad i enlighet med årsredovisningslagen och Bokföringsnämndens allmänna råd (BFNAR 2016:10) om årsredovisning i mindre företag.', space_after=8)
    if not manual.get('k2_previous_year', True):
        add_body_paragraph(doc, 'Årsredovisningen upprättas för första gången i enlighet med Bokföringsnämndens allmänna råd (BFNAR 2016:10) om årsredovisning i mindre företag, vilket kan innebära en bristande jämförbarhet mellan räkenskapsåret och det närmast föregående räkenskapsåret.', space_after=8)
    add_heading_clean(doc, 'Nyckeltalsdefinitioner', 3)
    add_heading_clean(doc, 'Nettoomsättning', 4)
    add_body_paragraph(doc, 'Rörelsens huvudintäkter, fakturerade kostnader, sidointäkter samt intäktskorrigeringar.', space_after=6)
    add_heading_clean(doc, 'Resultat efter finansiella poster', 4)
    add_body_paragraph(doc, 'Resultat efter finansiella intäkter och kostnader men före bokslutsdispositioner och skatter.', space_after=6)
    add_heading_clean(doc, 'Soliditet', 4)
    add_body_paragraph(doc, 'Justerat eget kapital (eget kapital och obeskattade reserver med avdrag för uppskjuten skatt) i procent av balansomslutningen.', space_after=8)

    add_heading_clean(doc, 'Not 2 - Medelantal anställda', 2)
    table = doc.add_table(rows=2, cols=2)
    configure_table(table, [9.5, 4.2], font_size=11.0, top_bottom_borders=True)
    set_cell_text_clean(table.cell(0,0), f"{raw['report_start']} - {raw['report_end']}", size=11.0)
    set_cell_text_clean(table.cell(0,1), '', size=11.0)
    set_cell_text_clean(table.cell(1,0), 'Medelantal anställda under året', size=11.0)
    set_cell_text_clean(table.cell(1,1), raw['avg_employees'], align=WD_ALIGN_PARAGRAPH.RIGHT, size=11.0)

    add_heading_clean(doc, 'Not 3 - Andelar i koncernföretag', 2)
    table = doc.add_table(rows=4, cols=2)
    configure_table(table, [9.5, 4.2], font_size=11.0, top_bottom_borders=True)
    rows = [
        (f"Anskaffningsvärden {raw['report_end']}", ''),
        ('Ingående anskaffningsvärden', format_kr(raw['balance']['shares_current'])),
        ('Utgående anskaffningsvärden', format_kr(raw['balance']['shares_current'])),
        ('Redovisat värde', format_kr(raw['balance']['shares_current'])),
    ]
    for i,(l,v) in enumerate(rows):
        set_cell_text_clean(table.cell(i,0), l, size=11.0)
        set_cell_text_clean(table.cell(i,1), v, align=WD_ALIGN_PARAGRAPH.RIGHT, size=11.0)

    if str(manual.get('significant_events_after_year_end', '')).strip():
        add_heading_clean(doc, 'Not 4 - Väsentliga händelser efter räkenskapsårets slut', 2)
        add_body_paragraph(doc, str(manual.get('significant_events_after_year_end', '')).strip(), space_after=6)
        next_note = 5
    else:
        next_note = 4

    group_interest_current = to_positive(get_account(sie.res, 0, 8423))
    group_interest_prior = to_positive(get_account(sie.res, -1, 8423))
    if group_interest_current or group_interest_prior:
        add_heading_clean(doc, f'Not {next_note} - Räntekostnader till koncernföretag', 2)
        add_body_paragraph(doc, f"Av årets räntekostnader och liknande resultatposter avser {format_kr(round_int_by_mode(group_interest_current, str(manual.get('rounding_mode', 'truncate'))))} kr ({format_kr(round_int_by_mode(group_interest_prior, str(manual.get('rounding_mode', 'truncate'))))} kr) skulder till koncernföretag.", space_after=6)


def add_signature_page(doc: Document, built: BuiltValues):
    raw = built.raw
    add_heading_clean(doc, 'Underskrifter', 1)
    add_body_paragraph(doc, f"Årsredovisning för {raw['company_name']}, {raw['org_number']} Avseende räkenskapsåret {raw['report_start']} - {raw['report_end']}", space_after=18)
    add_body_paragraph(doc, raw['board_city'], space_after=18)
    members = raw.get('board_members', []) or []
    cols = 2 if len(members) > 1 else 1
    rows = (len(members) + cols - 1) // cols if members else 1
    table = doc.add_table(rows=rows, cols=cols)
    configure_table(table, [6.4] * cols, font_size=11.0, top_bottom_borders=False)
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            if idx < len(members):
                member = members[idx]
                set_cell_text_clean(cell, '', size=11.0)
                p = cell.paragraphs[0]
                clear_paragraph(p)
                set_paragraph_base(p, space_before=0, space_after=0, line_spacing=1.0)
                add_text(p, str(member.get('name', '')).strip(), size=11.0)
                p = cell.add_paragraph()
                set_paragraph_base(p, space_before=0, space_after=0, line_spacing=1.0)
                add_text(p, str(member.get('title', '')).strip(), size=11.0)
                p = cell.add_paragraph()
                set_paragraph_base(p, space_before=0, space_after=0, line_spacing=1.0)
                add_text(p, str(member.get('date', raw.get('adoption_date', ''))).strip(), size=11.0)
            idx += 1


def add_management_page(doc: Document, built: BuiltValues, manual: Dict[str, Any]):
    raw = built.raw
    t0 = built.tables['table0']
    add_heading_clean(doc, 'Förvaltningsberättelse', 1)
    add_heading_clean(doc, 'Verksamheten', 2)
    add_heading_clean(doc, 'Allmänt om verksamheten', 3)
    add_body_paragraph(doc, f"{manual.get('business_description','').strip()} Företaget har sitt säte i {manual.get('registered_seat','').strip()}.", space_after=6)
    if str(manual.get('significant_events_during_year', '')).strip():
        add_heading_clean(doc, 'Väsentliga händelser under räkenskapsåret', 3)
        add_body_paragraph(doc, str(manual.get('significant_events_during_year', '')).strip(), space_after=6)
    if manual.get('has_own_shares'):
        add_heading_clean(doc, 'Egna aktier', 3)
        add_body_paragraph(doc, str(manual.get('own_shares_text', '')).strip(), space_after=6)
    if str(manual.get('net_sales_variation_comment', '')).strip():
        add_heading_clean(doc, 'Kommentar till flerårsöversikten', 3)
        add_body_paragraph(doc, str(manual.get('net_sales_variation_comment', '')).strip(), space_after=6)

    add_heading_clean(doc, 'Flerårsöversikt', 2)
    add_flerarsoversikt_table(doc, t0)
    add_spacer(doc, 8)
    add_heading_clean(doc, 'Förändringar i eget kapital', 2)
    add_equity_change_table(doc, t0)
    add_spacer(doc, 8)
    add_heading_clean(doc, 'Resultatdisposition', 2)
    add_result_disposition_table(doc, t0)


def generate_clean_docx(output_path: Path, built: BuiltValues, manual: Dict[str, Any], sie: SieData) -> None:
    doc = Document()
    set_document_defaults(doc)
    add_footer_clean(doc, built.raw['company_name'], built.raw['org_number'])

    add_cover_page(doc, built)
    doc.add_page_break()
    add_faststallelse_page(doc, built)
    doc.add_page_break()
    add_management_page(doc, built, manual)
    doc.add_page_break()
    add_two_year_statement_table(doc, 'Resultaträkning', built.raw['report_start'] + ' - ' + built.raw['report_end'], built.raw['prior_start'] + ' - ' + built.raw['prior_end'], built.tables['table1'], [(0,'Rörelsekostnader'),(4,'Finansiella poster'),(7,'Bokslutsdispositioner'),(11,'Skatter')], note_label='ja', font_size=10.5)
    doc.add_page_break()
    add_two_year_statement_table(doc, 'Balansräkning', built.raw['report_end'], built.raw['prior_end'], built.tables['table2'], [(0,'Tillgångar'),(0,'Anläggningstillgångar'),(0,'Finansiella anläggningstillgångar'),(3,'Omsättningstillgångar'),(3,'Kortfristiga fordringar'),(6,'Kassa och bank')], note_label='ja', font_size=10.5)
    doc.add_page_break()
    add_two_year_statement_table(doc, 'Balansräkning', built.raw['report_end'], built.raw['prior_end'], [[('Skulder till koncernföretag' if row[0]=='Skulder till koncernföretag_lång' else 'Skulder till koncernföretag' if row[0]=='Skulder till koncernföretag_kort' else row[0]), row[1], row[2]] for row in built.tables['table3']], [(0,'Eget kapital'),(0,'Bundet eget kapital'),(2,'Fritt eget kapital'),(6,'Obeskattade reserver'),(8,'Långfristiga skulder'),(11,'Kortfristiga skulder')], note_label='ja', font_size=10.5)
    doc.add_page_break()
    add_notes_page(doc, built, manual, sie)
    doc.add_page_break()
    add_signature_page(doc, built)
    doc.save(output_path)


def populate_docx(template_path: Optional[Path], output_path: Path, built: BuiltValues, manual: Dict[str, Any], sie: SieData, export_managed_template: Optional[Path] = None) -> None:
    generate_clean_docx(output_path, built, manual, sie)


def apply_overrides(values: Dict[str, Any], overrides: Dict[str, Any]) -> None:
    for path, override_value in overrides.items():
        target = values
        parts = path.split(".")
        for part in parts[:-1]:
            if part.isdigit() and isinstance(target, list):
                target = target[int(part)]
            else:
                if part not in target or not isinstance(target[part], (dict, list)):
                    raise KeyError(f"Kan inte hitta override-sökväg: {path}")
                target = target[part]
        last = parts[-1]
        if isinstance(target, list) and last.isdigit():
            target[int(last)] = override_value
        else:
            if last not in target:
                raise KeyError(f"Kan inte hitta override-sökväg: {path}")
            target[last] = override_value



def replace_runs(paragraph, text: str) -> None:
    text = text or ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)



def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        replace_runs(cell.paragraphs[0], text)
        for extra in cell.paragraphs[1:]:
            replace_runs(extra, "")
    else:
        cell.text = text



def parse_iso_date(value: Any) -> Optional[date]:
    if not value:
        return None
    try:
        return date.fromisoformat(str(value))
    except ValueError:
        return None


def add_months(base: date, months: int) -> date:
    year = base.year + (base.month - 1 + months) // 12
    month = (base.month - 1 + months) % 12 + 1
    day = min(base.day, [31, 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1])
    return date(year, month, day)


def analyze_net_sales_variation(manual: Dict[str, Any], built: BuiltValues) -> List[str]:
    current_year = int(str(manual.get("report_start", "2024"))[:4]) if str(manual.get("report_start", "2024"))[:4].isdigit() else 2024
    values = {
        current_year: int(built.raw["result"]["net_sales_current"]),
        current_year - 1: int(built.raw["result"]["net_sales_prior"]),
    }
    for year, payload in manual.get("flerarsoversikt", {}).items():
        if str(year).isdigit():
            values[int(year)] = int(payload.get("nettoomsattning", 0))
    affected = []
    years = sorted(values.keys(), reverse=True)
    for i in range(len(years) - 1):
        newer = years[i]
        older = years[i + 1]
        nv = values[newer]
        ov = values[older]
        if ov == 0 and nv == 0:
            continue
        if ov == 0 and nv != 0:
            affected.append(f"{older}->{newer}")
            continue
        change = abs(nv - ov) / abs(ov)
        if change > 0.30:
            affected.append(f"{older}->{newer}")
    return affected


def validate_rules(sie: SieData, manual: Dict[str, Any], built: BuiltValues) -> Dict[str, List[str]]:
    errors: List[str] = []
    warnings: List[str] = []
    infos: List[str] = []

    start = parse_iso_date(manual.get("report_start"))
    end = parse_iso_date(manual.get("report_end"))
    adoption = parse_iso_date(manual.get("adoption_date"))
    submission = parse_iso_date(manual.get("submission_date"))
    document_date = parse_iso_date(manual.get("document_date"))

    if built.raw["balance"]["assets_total_current"] != built.raw["balance"]["equity_and_liab_total_current"]:
        errors.append("Balansräkningen balanserar inte efter avrundning.")
    else:
        infos.append("Balansräkningen balanserar efter avrundning.")

    if manual.get("company_is_public"):
        errors.append("Publika aktiebolag får inte tillämpa K2.")
    if manual.get("parent_in_larger_group"):
        errors.append("Moderföretag i större koncern får inte tillämpa K2.")
    if manual.get("parent_in_smaller_group_prepares_consolidated"):
        errors.append("Moderföretag i mindre koncern som upprättar koncernredovisning får inte tillämpa K2.")
    if manual.get("foreign_branch"):
        errors.append("Företag med filial i utlandet under räkenskapsåret får inte tillämpa K2.")
    if manual.get("share_based_payments"):
        errors.append("Företag med aktierelaterade ersättningar får inte tillämpa K2.")
    if manual.get("compound_instruments"):
        errors.append("Företag med konvertibler eller liknande sammansatta finansiella instrument får inte tillämpa K2.")
    if manual.get("crypto_assets"):
        errors.append("Företag med kryptotillgångar får normalt inte tillämpa K2.")

    if start and start >= date(2024, 7, 1):
        infos.append("Räkenskapsåret omfattas av de nya K2-begränsningarna i BFNAR 2025:2 och av Bolagsverkets nya dateringskrav.")
        special_condition = bool(manual.get("deferred_tax_liability_significant")) or bool(manual.get("buildings_generate_75_pct_turnover"))
        thresholds = sum(1 for key in ["headcount_over_3_two_years", "balance_over_1_5m_two_years", "net_sales_over_3m_two_years"] if manual.get(key))
        if special_condition and thresholds >= 2:
            if manual.get("k2_previous_year"):
                warnings.append("Företaget träffas av K2-begränsningarna i punkt 1.1B, men undantaget för företag som normalt inte omfattas kan vara aktuellt. Manuell bedömning krävs.")
            else:
                errors.append("Företaget verkar inte få tillämpa K2 för detta räkenskapsår enligt BFNAR 2016:10 punkt 1.1B–1.1C.")
        elif special_condition and thresholds < 2:
            infos.append("Företaget har en 1.1B-omständighet men ligger inom undantaget när högst ett tröskelvillkor överskrids.")
    else:
        infos.append("Räkenskapsåret börjar före 2024-07-01, så de nya K2-begränsningarna i BFNAR 2025:2 gäller normalt inte detta år.")

    if end and adoption:
        latest_meeting = add_months(end, 6)
        if adoption > latest_meeting:
            errors.append(f"Årsstämman är senare än sex månader efter räkenskapsårets slut ({latest_meeting.isoformat()}).")
        else:
            infos.append("Årsstämmodatum ligger inom sexmånadersfristen.")
    elif end and not adoption:
        warnings.append("Stämmodatum saknas, så scriptet kan inte kontrollera sexmånadersfristen.")

    if submission and adoption:
        filing_deadline = add_months(adoption, 1)
        if submission > filing_deadline:
            warnings.append(f"Planerat inlämningsdatum är senare än en månad efter årsstämman ({filing_deadline.isoformat()}).")
        else:
            infos.append("Planerat inlämningsdatum ligger inom en månad efter årsstämman.")
    elif adoption and not submission:
        warnings.append("Planerat inlämningsdatum saknas, så scriptet kan inte kontrollera månadsfristen till Bolagsverket.")

    if submission and end:
        fee_deadline = add_months(end, 7)
        if submission > fee_deadline:
            warnings.append(f"Planerat inlämningsdatum är senare än sju månader efter räkenskapsårets slut ({fee_deadline.isoformat()}) och riskerar förseningsavgift.")
        else:
            infos.append("Planerat inlämningsdatum ligger inom sju månader från räkenskapsårets slut.")

    if start and start >= date(2024, 7, 1) and not document_date:
        warnings.append("Dateringsdatum för själva årsredovisningen saknas trots att nya dateringsregeln gäller.")

    if manual.get("has_auditor"):
        infos.append("Bolaget har revisor. Revisionsberättelse måste följa med vid inlämning.")

    report_mode = str(manual.get("report_mode", "paper"))
    if report_mode == "digital":
        warnings.append("Det här scriptet skapar en DOCX-fil. Digital inlämning till Bolagsverket kräver iXBRL/XBRL via program eller tjänst som stöder det.")
    else:
        infos.append("Scriptet är anpassat för arbetskopia/pappersspår. Fastställelseintyget ska inte ligga på separat papper.")

    if manual.get("has_own_shares") and not str(manual.get("own_shares_text", "")).strip():
        errors.append("Bolaget har egna aktier, men text för upplysningen saknas.")
    elif manual.get("has_own_shares"):
        infos.append("Upplysning om egna aktier kommer att infogas automatiskt i förvaltningsberättelsen.")

    if manual.get("significant_events_during_year"):
        infos.append("Sektion för väsentliga händelser under räkenskapsåret kommer att infogas automatiskt.")
    if manual.get("significant_events_after_year_end"):
        infos.append("Not om väsentliga händelser efter räkenskapsårets slut kommer att infogas automatiskt.")

    if abs(sum_accounts(sie.res, 0, [8423])) > ZERO or abs(sum_accounts(sie.res, -1, [8423])) > ZERO:
        infos.append("Scriptet lägger in en separat not om räntekostnader till koncernföretag när konto 8423 används.")

    varied_years = analyze_net_sales_variation(manual, built)
    if varied_years and not str(manual.get("net_sales_variation_comment", "")).strip():
        warnings.append("Nettoomsättningen varierar mer än 30 procent mellan år i flerårsöversikten (" + ", ".join(varied_years) + "). K2 kräver kommentar.")
    elif varied_years:
        infos.append("Kommentar om variation i nettoomsättning finns angiven.")

    infos.append("Förvaltningsberättelsens K2-rubriker omfattar i scriptet Verksamheten, Flerårsöversikt, Förändringar i eget kapital och Resultatdisposition, med dynamiska tillägg för väsentliga händelser, egna aktier och kommentar till flerårsöversikten när det behövs.")
    infos.append("Not 1 om redovisningsprinciper, not om medelantal anställda, not om andelar i koncernföretag och extra noter vid behov är inbyggda i scriptets mallfyllning.")
    return {"errors": errors, "warnings": warnings, "infos": infos}


def build_report(sie: SieData, built: BuiltValues, manual: Dict[str, Any], validations: Dict[str, List[str]]) -> str:
    lines = [
        "# Årsredovisning från SIE – körningsrapport",
        "",
        f"Bolagsnamn i SIE: **{sie.company_name}**",
        f"Organisationsnummer i SIE: **{sie.org_number}**",
        f"Räkenskapsår i SIE: **{sie.current_start} - {sie.current_end}**",
        "",
        "## Regelmotor – sammanfattning",
        "",
        f"- Inlämningssätt: **{manual.get('report_mode', 'paper')}**",
        f"- Revisor: **{'ja' if manual.get('has_auditor') else 'nej'}**",
        f"- Dateringsdatum: **{manual.get('document_date', '') or 'saknas'}**",
        f"- Stämmodatum: **{manual.get('adoption_date', '') or 'saknas'}**",
        f"- Planerat inlämningsdatum: **{manual.get('submission_date', '') or 'saknas'}**",
        "",
        "## Regelfel som måste hanteras",
        "",
    ]
    if validations['errors']:
        lines.extend(f"- {item}" for item in validations['errors'])
    else:
        lines.append("- Inga blockerande regelfel upptäckta i den automatiska kontrollen.")
    lines.extend(["", "## Varningar och manuella kontroller", ""])
    if validations['warnings']:
        lines.extend(f"- {item}" for item in validations['warnings'])
    else:
        lines.append("- Inga varningar i den automatiska kontrollen.")
    lines.extend(["", "## Informationspunkter", ""])
    lines.extend(f"- {item}" for item in validations['infos'])
    lines.extend([
        "",
        "## Avrundning",
        "",
        f"Scriptet använder avrundningsläge **{built.raw['rounding_mode']}** och balanserar därefter raderna till hela kronor.",
        f"- Visad total 2024 efter balanserad avrundning: **{format_kr(built.raw['balance']['assets_total_current'])}**",
        f"- Justeringssteg tillgångssidan 2024: **{built.raw['rounding']['assets_current_adjustment_steps']}**",
        f"- Justeringssteg EK/skuldsidan 2024: **{built.raw['rounding']['liabs_current_adjustment_steps']}**",
        f"- Justeringssteg resultaträkning 2024: **{built.raw['rounding']['result_current_adjustment_steps']}**",
        "",
        "## Viktiga exempelvärden",
        "",
        f"- Resultat efter finansiella poster 2024: **{format_kr(built.raw['result']['after_fin_current'])}**",
        f"- Årets resultat 2024: **{format_kr(built.raw['result']['year_current'])}**",
        f"- Summa tillgångar 2024: **{format_kr(built.raw['balance']['assets_total_current'])}**",
        f"- Summa eget kapital och skulder 2024: **{format_kr(built.raw['balance']['equity_and_liab_total_current'])}**",
        "",
        "## Användning",
        "",
        "Interaktiv körning:",
        "```bash",
        "python generate_arsredovisning_from_sie_v6.py",
        "```",
        "",
        "Strikt körning som stoppar vid regelbrott:",
        "```bash",
        "python generate_arsredovisning_from_sie_v6.py --sie bokforing.se --template mall.docx --output arsredovisning.docx --manual manual.json",
        "```",
        "",
    ])
    if manual.get("overrides"):
        lines.append("## Aktiva overrides")
        lines.append("")
        lines.extend(f"- `{key}` = `{value}`" for key, value in manual["overrides"].items())
        lines.append("")
    return "\n".join(lines)



def normalize_user_path(raw: str) -> Path:
    raw = raw.strip().strip('"').strip("'")
    return Path(raw).expanduser()


def prompt_existing_file(label: str, default: str = "") -> Path:
    while True:
        raw = prompt_text(label, default)
        path = normalize_user_path(raw)
        if path.exists() and path.is_file():
            return path
        print("Filen hittades inte. Kontrollera sökvägen och försök igen.")



def sanitize_filename_component(value: str, fallback: str) -> str:
    value = re.sub(r'[<>:"/\|?*]', '', (value or '').strip())
    value = re.sub(r'\s+', ' ', value).strip(' .')
    return value[:120] if value else fallback



def guess_template_path(script_dir: Path) -> Optional[Path]:
    candidates = [
        script_dir / 'managed_template_v4.docx',
        script_dir / 'managed_template_local.docx',
        script_dir / 'Skånmyr Nilsson årsredovisning 2024.docx',
    ]
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    docx_files = sorted(path for path in script_dir.glob('*.docx') if path.is_file())
    return docx_files[0] if docx_files else None



def find_first_input_file_in_testing(script_dir: Path) -> Path:
    testing_dir = script_dir / "testing"
    if not testing_dir.exists() or not testing_dir.is_dir():
        raise FileNotFoundError(
            f"Mappen '{testing_dir}' finns inte. Skapa en undermapp som heter testing bredvid scriptet och lägg SIE-filen där."
        )

    files = sorted(path for path in testing_dir.iterdir() if path.is_file())
    if not files:
        raise FileNotFoundError(
            f"Mappen '{testing_dir}' är tom. Lägg en SIE-fil där och kör scriptet igen."
        )

    preferred = [path for path in files if path.suffix.lower() in {'.se', '.sie', '.si'}]
    return preferred[0] if preferred else files[0]


def derive_output_paths(sie_path: Path, manual: Dict[str, Any]) -> Dict[str, Path]:
    company_name = sanitize_filename_component(str(manual.get('company_name', '')), 'Bolag')
    report_start = str(manual.get('report_start', ''))
    year = report_start[:4] if report_start[:4].isdigit() else 'okänt-år'
    base_name = f"{company_name} årsredovisning {year}"
    return {
        'output_docx': sie_path.parent / f"{base_name}.docx",
        'manual_json': sie_path.parent / f"{base_name} underlag.json",
        'report_md': sie_path.parent / f"{base_name} kontrollrapport.md",
    }



def run_guided_mode() -> None:
    script_dir = Path(__file__).resolve().parent
    print("Årsredovisning från SIE\n")

    try:
        sie_path = find_first_input_file_in_testing(script_dir)
    except FileNotFoundError as exc:
        raise SystemExit(str(exc))

    print(f"SIE-fil som används: {sie_path}")
    print("Layoutmotor: ren v7-layout (Times New Roman, städade tabeller, egen sidfot)")

    sie = parse_sie(sie_path)
    manual = ensure_manual_data(sie, {}, interactive=True)
    built = build_values(sie, manual)
    validations = validate_rules(sie, manual, built)

    paths = derive_output_paths(sie_path, manual)
    paths['report_md'].write_text(build_report(sie, built, manual, validations), encoding='utf-8')

    if validations["errors"]:
        print("\nScriptet stoppade eftersom det bara skapar K2-årsredovisning eller för att andra blockerande fel hittades.\n")
        for item in validations["errors"]:
            print(f"- {item}")
        print(f"\nKontrollrapport: {paths['report_md']}")
        raise SystemExit(1)

    populate_docx(None, paths['output_docx'], built, manual, sie)
    write_json(paths['manual_json'], manual)

    print("\nKlart.")
    print(f"Årsredovisning: {paths['output_docx']}")
    print(f"Manuellt underlag: {paths['manual_json']}")
    print(f"Kontrollrapport: {paths['report_md']}")


def main() -> None:
    if len(sys.argv) == 1:
        run_guided_mode()
        return

    parser = argparse.ArgumentParser(description="Läser första SIE-filen i testing-mappen eller en angiven SIE-fil och skapar en ren K2-årsredovisning i DOCX-format.")
    parser.add_argument("--sie", type=Path, help="Valfri sökväg till SIE-fil. Om utelämnad används första filen i testing-mappen.")
    parser.add_argument("--output", type=Path, help="Valfri utdata-DOCX")
    parser.add_argument("--manual", type=Path, help="JSON-fil med manuella uppgifter")
    parser.add_argument("--save-manual", type=Path, help="Spara den slutliga manuella JSON-filen hit")
    parser.add_argument("--report", type=Path, help="Valfri rapport i Markdown-format")
    parser.add_argument("--interactive", action="store_true", help="Fråga efter manuella uppgifter i terminalen")
    parser.add_argument("--rounding-mode", choices=["truncate", "half_up"], help="Tvinga ett visst avrundningsläge")
    args = parser.parse_args()

    script_dir = Path(__file__).resolve().parent
    sie_path = args.sie or find_first_input_file_in_testing(script_dir)

    sie = parse_sie(sie_path)
    manual = read_json(args.manual) if args.manual and args.manual.exists() else {}
    if args.rounding_mode:
        manual["rounding_mode"] = args.rounding_mode
    manual = ensure_manual_data(sie, manual, interactive=args.interactive or not args.manual)
    built = build_values(sie, manual)
    validations = validate_rules(sie, manual, built)

    paths = derive_output_paths(sie_path, manual)
    output_path = args.output or paths['output_docx']
    report_path = args.report or paths['report_md']
    report_path.write_text(build_report(sie, built, manual, validations), encoding="utf-8")

    if validations["errors"]:
        raise SystemExit("Scriptet stoppade eftersom det bara skapar K2-årsredovisning eller för att andra blockerande fel hittades.")

    populate_docx(None, output_path, built, manual, sie)

    if args.save_manual:
        write_json(args.save_manual, manual)
    else:
        write_json(paths['manual_json'], manual)

    print(f"Klart: {output_path}")
    print(f"Rapport: {report_path}")


if __name__ == "__main__":
    main()
